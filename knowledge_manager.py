import json
import os
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, urljoin
import io
import logging
from typing import List, Dict, Any
import asyncio
import aiohttp

# Optional imports - will be handled in the processing functions
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class KnowledgeManager:
    def __init__(self, knowledge_file="knowledge_base.json"):
        self.knowledge_file = knowledge_file
        self.knowledge_base = self.load_knowledge_base()
    
    def load_knowledge_base(self) -> Dict[str, Any]:
        """Load existing knowledge base or create new one"""
        try:
            if os.path.exists(self.knowledge_file):
                with open(self.knowledge_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            else:
                return self.create_default_knowledge_base()
        except Exception as e:
            logger.error(f"Error loading knowledge base: {e}")
            return self.create_default_knowledge_base()
    
    def create_default_knowledge_base(self) -> Dict[str, Any]:
        """Create default knowledge base structure"""
        return {
            "company_info": {},
            "services": {},
            "scraped_urls": {},
            "uploaded_documents": {},
            "last_updated": "",
            "sources": []
        }
    
    def save_knowledge_base(self):
        """Save knowledge base to file"""
        try:
            with open(self.knowledge_file, 'w', encoding='utf-8') as f:
                json.dump(self.knowledge_base, f, indent=2, ensure_ascii=False)
            logger.info("Knowledge base saved successfully")
        except Exception as e:
            logger.error(f"Error saving knowledge base: {e}")
    
    async def scrape_url(self, url: str, max_depth: int = 1) -> Dict[str, Any]:
        """Scrape content from a URL"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    if response.status == 200:
                        html = await response.text()
                        soup = BeautifulSoup(html, 'html.parser')
                        
                        # Extract relevant content
                        content = {
                            "url": url,
                            "title": soup.find('title').text.strip() if soup.find('title') else "No title",
                            "meta_description": "",
                            "headings": [],
                            "paragraphs": [],
                            "links": [],
                            "scraped_at": ""
                        }
                        
                        # Meta description
                        meta_desc = soup.find('meta', attrs={'name': 'description'})
                        if meta_desc:
                            content["meta_description"] = meta_desc.get('content', '')
                        
                        # Headings
                        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                            content["headings"].append({
                                "level": heading.name,
                                "text": heading.text.strip()
                            })
                        
                        # Paragraphs
                        for p in soup.find_all('p'):
                            text = p.text.strip()
                            if text and len(text) > 20:  # Filter out short/empty paragraphs
                                content["paragraphs"].append(text)
                        
                        # Internal links (for potential further scraping)
                        base_domain = urlparse(url).netloc
                        for link in soup.find_all('a', href=True):
                            href = link['href']
                            if href.startswith('/'):
                                href = urljoin(url, href)
                            if base_domain in href:
                                content["links"].append({
                                    "url": href,
                                    "text": link.text.strip()
                                })
                        
                        # Store in knowledge base
                        self.knowledge_base["scraped_urls"][url] = content
                        self.knowledge_base["sources"].append({
                            "type": "url",
                            "source": url,
                            "title": content["title"]
                        })
                        
                        logger.info(f"Successfully scraped: {url}")
                        return content
                    else:
                        logger.error(f"Failed to scrape {url}: HTTP {response.status}")
                        return {}
        except Exception as e:
            logger.error(f"Error scraping {url}: {e}")
            return {}
    
    def process_pdf_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF document and extract text with fallback handling"""
        try:
            # Check if PyPDF2 is available
            if PyPDF2 is None:
                logger.error("PyPDF2 not available - PDF processing disabled")
                return {
                    "filename": filename,
                    "type": "pdf",
                    "content": f"PDF document '{filename}' uploaded but text extraction not available. Please install PyPDF2 for full functionality.",
                    "processed_at": "",
                    "error": "PyPDF2 not installed"
                }
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_content = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                try:
                    text_content.append(page.extract_text())
                except Exception as e:
                    text_content.append(f"[Error extracting page {page_num + 1}: {str(e)}]")
            
            doc_info = {
                "filename": filename,
                "type": "pdf",
                "pages": len(pdf_reader.pages),
                "content": "\n".join(text_content) if text_content else "No text content extracted",
                "processed_at": "",
                "summary": ""
            }
            
            # Store in knowledge base
            self.knowledge_base["uploaded_documents"][filename] = doc_info
            self.knowledge_base["sources"].append({
                "type": "document",
                "source": filename,
                "title": filename
            })
            
            logger.info(f"Successfully processed PDF: {filename}")
            return doc_info
            
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {e}")
            # Return a basic record even if processing fails
            return {
                "filename": filename,
                "type": "pdf",
                "content": f"PDF document '{filename}' uploaded but processing failed: {str(e)}",
                "processed_at": "",
                "error": str(e)
            }
    
    def process_docx_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process Word document and extract text with fallback handling"""
        try:
            # Check if python-docx is available
            if Document is None:
                logger.error("python-docx not available - Word document processing disabled")
                return {
                    "filename": filename,
                    "type": "docx",
                    "content": f"Word document '{filename}' uploaded but text extraction not available. Please install python-docx for full functionality.",
                    "processed_at": "",
                    "error": "python-docx not installed"
                }
            
            doc = Document(io.BytesIO(file_content))
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            doc_info = {
                "filename": filename,
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "content": "\n".join(text_content) if text_content else "No text content found",
                "processed_at": "",
                "summary": ""
            }
            
            # Store in knowledge base
            self.knowledge_base["uploaded_documents"][filename] = doc_info
            self.knowledge_base["sources"].append({
                "type": "document",
                "source": filename,
                "title": filename
            })
            
            logger.info(f"Successfully processed DOCX: {filename}")
            return doc_info
            
        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {e}")
            # Return a basic record even if processing fails
            return {
                "filename": filename,
                "type": "docx",
                "content": f"Word document '{filename}' uploaded but processing failed: {str(e)}",
                "processed_at": "",
                "error": str(e)
            }
    
    def search_knowledge(self, query: str, max_results: int = 5) -> List[Dict[str, Any]]:
        """Search through knowledge base for relevant information"""
        results = []
        query_lower = query.lower()
        
        # Search scraped URLs
        for url, content in self.knowledge_base.get("scraped_urls", {}).items():
            relevance_score = 0
            
            # Check title
            if query_lower in content.get("title", "").lower():
                relevance_score += 3
            
            # Check meta description
            if query_lower in content.get("meta_description", "").lower():
                relevance_score += 2
            
            # Check headings
            for heading in content.get("headings", []):
                if query_lower in heading.get("text", "").lower():
                    relevance_score += 2
            
            # Check paragraphs
            for paragraph in content.get("paragraphs", []):
                if query_lower in paragraph.lower():
                    relevance_score += 1
            
            if relevance_score > 0:
                results.append({
                    "type": "url",
                    "source": url,
                    "title": content.get("title", ""),
                    "relevance": relevance_score,
                    "content": content
                })
        
        # Search uploaded documents
        for filename, doc_info in self.knowledge_base.get("uploaded_documents", {}).items():
            relevance_score = 0
            
            # Check filename
            if query_lower in filename.lower():
                relevance_score += 2
            
            # Check content
            if query_lower in doc_info.get("content", "").lower():
                relevance_score += doc_info.get("content", "").lower().count(query_lower)
            
            if relevance_score > 0:
                results.append({
                    "type": "document",
                    "source": filename,
                    "title": filename,
                    "relevance": relevance_score,
                    "content": doc_info
                })
        
        # Sort by relevance and return top results
        results.sort(key=lambda x: x["relevance"], reverse=True)
        return results[:max_results]
    
    def get_context_for_query(self, query: str) -> str:
        """Get relevant context from knowledge base for AI query"""
        search_results = self.search_knowledge(query)
        
        if not search_results:
            return "No specific information found in knowledge base."
        
        context_parts = []
        context_parts.append("Relevant information from STAFFVIRTUAL knowledge base:")
        
        for result in search_results:
            if result["type"] == "url":
                content = result["content"]
                context_parts.append(f"\nFrom {result['source']}:")
                context_parts.append(f"Title: {content.get('title', '')}")
                if content.get("meta_description"):
                    context_parts.append(f"Description: {content['meta_description']}")
                
                # Add relevant paragraphs
                relevant_paragraphs = []
                query_lower = query.lower()
                for paragraph in content.get("paragraphs", []):
                    if query_lower in paragraph.lower():
                        relevant_paragraphs.append(paragraph)
                
                if relevant_paragraphs:
                    context_parts.append("Relevant content:")
                    context_parts.extend(relevant_paragraphs[:3])  # Limit to 3 paragraphs
            
            elif result["type"] == "document":
                doc_info = result["content"]
                context_parts.append(f"\nFrom document {result['source']}:")
                
                # Add relevant content snippets
                content = doc_info.get("content", "")
                query_lower = query.lower()
                sentences = content.split('.')
                relevant_sentences = [s.strip() for s in sentences if query_lower in s.lower()]
                
                if relevant_sentences:
                    context_parts.append("Relevant content:")
                    context_parts.extend(relevant_sentences[:5])  # Limit to 5 sentences
        
        return "\n".join(context_parts)
    
    def add_url_to_scrape_list(self, url: str):
        """Add URL to list of sources to scrape"""
        if "urls_to_scrape" not in self.knowledge_base:
            self.knowledge_base["urls_to_scrape"] = []
        
        if url not in self.knowledge_base["urls_to_scrape"]:
            self.knowledge_base["urls_to_scrape"].append(url)
            self.save_knowledge_base()
    
    def get_knowledge_summary(self) -> Dict[str, Any]:
        """Get summary of current knowledge base"""
        return {
            "scraped_urls": len(self.knowledge_base.get("scraped_urls", {})),
            "uploaded_documents": len(self.knowledge_base.get("uploaded_documents", {})),
            "total_sources": len(self.knowledge_base.get("sources", [])),
            "last_updated": self.knowledge_base.get("last_updated", "Never")
        }
